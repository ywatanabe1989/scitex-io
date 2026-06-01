#!/usr/bin/env python3
# File: ./scitex_repo/tests/scitex/io/_load_modules/test__docx.py

"""Real-collaborator tests for ``_load_docx``.

The original suite patched ``docx.Document`` and asserted on call
shape — call-record theater. We replace it with real python-docx
files written into ``tmp_path``.
"""

from __future__ import annotations

import pytest

# Required for scitex.io module
pytest.importorskip("h5py")
pytest.importorskip("zarr")
# python-docx (`from docx import Document`) is required by the DOCX loader.
pytest.importorskip("docx")

from docx import Document  # noqa: E402


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _write_docx(path, paragraphs):
    """Write a real .docx file with the given paragraphs to ``path``."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Extension validation
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "invalid_path",
    ["file.txt", "document.doc", "text.pdf", "data.xlsx"],
)
def test_load_docx_rejects_non_docx_extension(invalid_path):
    """Reject paths whose extension is not ``.docx``."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    # Act
    ctx = pytest.raises(ValueError, match="File must have .docx extension")
    # Assert
    with ctx:
        _load_docx(invalid_path)


# ---------------------------------------------------------------------------
# Concatenation behaviour
# ---------------------------------------------------------------------------


def test_load_docx_concatenates_paragraphs_into_single_string(tmp_path):
    """Paragraphs are concatenated (no separator) into one string."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    paragraphs = [
        "This is the first paragraph.",
        "This is the second paragraph.",
        "Final paragraph here.",
    ]
    docx_path = _write_docx(tmp_path / "test_document.docx", paragraphs)
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert result == "".join(paragraphs)


def test_load_docx_returns_str(tmp_path):
    """``_load_docx`` returns a ``str``."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    docx_path = _write_docx(tmp_path / "doc.docx", ["hello"])
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert isinstance(result, str)


# ---------------------------------------------------------------------------
# Empty document
# ---------------------------------------------------------------------------


def test_load_docx_empty_document_returns_empty_string(tmp_path):
    """A document with no paragraphs yields an empty string."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    docx_path = tmp_path / "empty.docx"
    Document().save(str(docx_path))
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert result == ""


def test_load_docx_empty_document_returns_str(tmp_path):
    """Empty document still returns ``str``."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    docx_path = tmp_path / "empty.docx"
    Document().save(str(docx_path))
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert isinstance(result, str)


# ---------------------------------------------------------------------------
# Single paragraph
# ---------------------------------------------------------------------------


def test_load_docx_single_paragraph_text_matches(tmp_path):
    """A single-paragraph document returns that paragraph verbatim."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    docx_path = _write_docx(tmp_path / "one.docx", ["Only paragraph."])
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert result == "Only paragraph."


# ---------------------------------------------------------------------------
# Unicode content
# ---------------------------------------------------------------------------


def test_load_docx_preserves_unicode_content(tmp_path):
    """Unicode paragraphs survive the round-trip."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    paragraphs = ["日本語のテキスト", "Ελληνικά", "Русский"]
    docx_path = _write_docx(tmp_path / "unicode.docx", paragraphs)
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert result == "".join(paragraphs)


# ---------------------------------------------------------------------------
# Missing file
# ---------------------------------------------------------------------------


def test_load_docx_missing_file_raises(tmp_path):
    """A non-existent .docx path raises an error (file-not-found or package)."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    # Act
    ctx = pytest.raises(Exception)
    # Assert
    with ctx:
        _load_docx(str(tmp_path / "missing.docx"))


# ---------------------------------------------------------------------------
# Large document
# ---------------------------------------------------------------------------


def test_load_docx_large_document_concatenation_length(tmp_path):
    """A 200-paragraph document yields the concatenated length."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    paragraphs = [f"Paragraph {i}." for i in range(200)]
    docx_path = _write_docx(tmp_path / "large.docx", paragraphs)
    expected = "".join(paragraphs)
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert len(result) == len(expected)


def test_load_docx_large_document_concatenation_content(tmp_path):
    """A 200-paragraph document yields the concatenated content."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    paragraphs = [f"Paragraph {i}." for i in range(200)]
    docx_path = _write_docx(tmp_path / "large.docx", paragraphs)
    expected = "".join(paragraphs)
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert result == expected


# ---------------------------------------------------------------------------
# Whitespace handling
# ---------------------------------------------------------------------------


def test_load_docx_preserves_whitespace_in_paragraphs(tmp_path):
    """Inner whitespace within paragraphs survives."""
    # Arrange
    from scitex_io._load_modules._docx import _load_docx
    paragraphs = ["  leading", "trailing  ", "in  ner"]
    docx_path = _write_docx(tmp_path / "ws.docx", paragraphs)
    # Act
    result = _load_docx(str(docx_path))
    # Assert
    assert result == "".join(paragraphs)


if __name__ == "__main__":
    import os
    pytest.main([os.path.abspath(__file__)])
