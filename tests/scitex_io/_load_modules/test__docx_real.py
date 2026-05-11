#!/usr/bin/env python3
"""Real-fixture tests for scitex_io._load_modules._docx._load_docx."""

import pytest

pytest.importorskip("docx")
from docx import Document

from scitex_io._load_modules._docx import _load_docx


def test_empty_docx(tmp_path):
    p = tmp_path / "empty.docx"
    Document().save(p)
    assert _load_docx(str(p)) == ""


def test_single_paragraph(tmp_path):
    p = tmp_path / "single.docx"
    doc = Document()
    doc.add_paragraph("hello world")
    doc.save(p)
    assert _load_docx(str(p)) == "hello world"


def test_multi_paragraphs(tmp_path):
    p = tmp_path / "multi.docx"
    doc = Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("beta")
    doc.add_paragraph("gamma")
    doc.save(p)
    out = _load_docx(str(p))
    assert "alpha" in out and "beta" in out and "gamma" in out


def test_wrong_extension_raises(tmp_path):
    with pytest.raises(ValueError, match="File must have .docx extension"):
        _load_docx(str(tmp_path / "x.txt"))


def test_nonexistent_file(tmp_path):
    # python-docx raises PackageNotFoundError for missing file
    from docx.opc.exceptions import PackageNotFoundError

    with pytest.raises(PackageNotFoundError):
        _load_docx(str(tmp_path / "missing.docx"))
