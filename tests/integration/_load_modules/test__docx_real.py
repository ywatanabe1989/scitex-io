#!/usr/bin/env python3
"""Real-fixture tests for scitex_io._load_modules._docx._load_docx."""

import pytest

pytest.importorskip("docx")
from docx import Document

from scitex_io._load_modules._docx import _load_docx


def test_empty_docx_load_docx_str_p(tmp_path):
    # Arrange
    # Arrange
    p = tmp_path / "empty.docx"
    # Act
    # Act
    Document().save(p)
    # Assert
    # Assert
    assert _load_docx(str(p)) == ""


def test_single_paragraph_load_docx_str_p_hello_world(tmp_path):
    # Arrange
    # Arrange
    p = tmp_path / "single.docx"
    doc = Document()
    doc.add_paragraph("hello world")
    # Act
    # Act
    doc.save(p)
    # Assert
    # Assert
    assert _load_docx(str(p)) == "hello world"


def test_multi_paragraphs_alpha_in_out_and_beta_in_out_and_gamma_in_out(tmp_path):
    # Arrange
    # Arrange
    p = tmp_path / "multi.docx"
    doc = Document()
    doc.add_paragraph("alpha")
    doc.add_paragraph("beta")
    doc.add_paragraph("gamma")
    doc.save(p)
    # Act
    # Act
    out = _load_docx(str(p))
    # Assert
    # Assert
    assert "alpha" in out and "beta" in out and "gamma" in out


def test_wrong_extension_raises(tmp_path):
    # Arrange
    # Act
    # Assert
    # Arrange
    # Act
    # Assert
    with pytest.raises(ValueError, match="File must have .docx extension"):
        _load_docx(str(tmp_path / "x.txt"))


def test_nonexistent_file_raises_packagenotfounderror(tmp_path):
    # python-docx raises PackageNotFoundError for missing file
    # Arrange
    # Act
    # Arrange
    # Act
    from docx.opc.exceptions import PackageNotFoundError

    # Assert
    # Assert
    with pytest.raises(PackageNotFoundError):
        _load_docx(str(tmp_path / "missing.docx"))
